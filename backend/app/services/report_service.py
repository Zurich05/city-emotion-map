from io import BytesIO
from typing import TYPE_CHECKING

from app.services.hotspot_service import HotspotService
from app.services.statistics_service import StatisticsService

if TYPE_CHECKING:
    from sqlalchemy.orm import Session


class ReportService:
    def __init__(self, db: "Session"):
        self.db = db

    def generate(self) -> dict:
        overview = StatisticsService(self.db).overview()
        hotspots = HotspotService(self.db).list_hotspots(5)
        lines = [
            "城市情绪地图系统摘要报告",
            f"当前已分析样本 {overview['total_count']} 条，平均压力指数 {overview['avg_stress_score']}，高风险区域 {overview['high_risk_count']} 个。",
            f"积极比例 {overview['positive_ratio']}，消极比例 {overview['negative_ratio']}，中性比例 {overview['neutral_ratio']}。",
        ]
        for item in hotspots:
            lines.append(f"{item['district']} - {item['location_name']}：{item['main_reason']} {item['suggestion']}")
        return {"summary": "\n".join(lines), "overview": overview, "hotspots": hotspots}

    def export_pdf(self) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfgen import canvas

        report = self.generate()
        buffer = BytesIO()
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        y = height - 48
        pdf.setTitle("city-emotion-report")
        pdf.setFont("STSong-Light", 12)
        for line in report["summary"].splitlines():
            pdf.drawString(48, y, line[:72])
            y -= 22
            if y < 48:
                pdf.showPage()
                pdf.setFont("STSong-Light", 12)
                y = height - 48
        pdf.save()
        return buffer.getvalue()

    def export_docx(self) -> bytes:
        from docx import Document

        report = self.generate()
        document = Document()
        document.add_heading("城市情绪地图系统摘要报告", level=1)
        for line in report["summary"].splitlines()[1:]:
            document.add_paragraph(line)
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
